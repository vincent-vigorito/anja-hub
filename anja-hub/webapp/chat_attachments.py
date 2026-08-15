"""chat_attachments.py — handle file attachments per chat (Fase 24).

Pipeline:
  client uploads multipart → classify MIME → extract text (text-based) o base64 (image) →
  return {file_id, type, mime, preview, extracted_text|image_b64}.

Storage: <webapp>/uploads/<conv_id>/<uuid>-<filename>
Cleanup: optional retention policy (manual or daily routine).

Estrattori (lazy import, fallback graceful):
  - PDF: pymupdf (fitz) → pdfplumber → PyPDF2
  - DOCX: python-docx
  - XLSX: openpyxl (riusa anja_office deps)
  - TXT/MD/CSV/JSON/code: raw read (size-capped)
"""

from __future__ import annotations

import base64
import json
import mimetypes
import re
import secrets
import shutil
from datetime import datetime
from pathlib import Path
from typing import Optional


# Max sizes per type (in bytes)
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
MAX_TEXT_FILE_SIZE = 5 * 1024 * 1024  # 5MB (raw text)
MAX_EXTRACTED_CHARS = 100_000  # 100k chars max nel prompt
MAX_IMAGE_SIZE = 20 * 1024 * 1024  # 20MB

# MIME → category
TEXT_MIMES = {"text/plain", "text/markdown", "text/csv", "application/json", "text/x-python", "text/javascript", "text/html", "text/xml"}
TEXT_EXTS = {".txt", ".md", ".csv", ".json", ".py", ".js", ".ts", ".jsx", ".tsx", ".html", ".xml", ".yaml", ".yml", ".toml", ".sh", ".go", ".rs", ".java", ".c", ".cpp", ".h", ".hpp", ".sql", ".css", ".log", ".env"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
XLSX_EXTS = {".xlsx", ".xls"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
AUDIO_EXTS = {".mp3", ".wav", ".ogg", ".m4a", ".flac"}


def uploads_root(webapp_dir: Path) -> Path:
    d = webapp_dir / "uploads"
    d.mkdir(parents=True, exist_ok=True)
    return d


def conv_uploads_dir(webapp_dir: Path, conv_id: str) -> Path:
    safe = re.sub(r"[^a-zA-Z0-9_-]+", "_", conv_id or "default")
    d = uploads_root(webapp_dir) / safe
    d.mkdir(parents=True, exist_ok=True)
    return d


def classify(filename: str, mime: Optional[str] = None) -> str:
    """Ritorna categoria: 'image' | 'pdf' | 'docx' | 'xlsx' | 'text' | 'audio' | 'binary'."""
    ext = Path(filename).suffix.lower()
    if ext in IMAGE_EXTS:
        return "image"
    if ext in PDF_EXTS:
        return "pdf"
    if ext in DOCX_EXTS:
        return "docx"
    if ext in XLSX_EXTS:
        return "xlsx"
    if ext in AUDIO_EXTS:
        return "audio"
    if ext in TEXT_EXTS:
        return "text"
    if mime and mime in TEXT_MIMES:
        return "text"
    if mime and mime.startswith("text/"):
        return "text"
    if mime and mime.startswith("image/"):
        return "image"
    return "binary"


# ============================================================
# Extractors (lazy import, graceful fallback)
# ============================================================

def _extract_pdf(path: Path) -> tuple[str, Optional[str]]:
    """Return (text, error). Try pymupdf → pdfplumber → PyPDF2."""
    # pymupdf (best quality)
    try:
        import fitz  # pymupdf
        doc = fitz.open(str(path))
        chunks = []
        for page in doc:
            chunks.append(page.get_text())
        doc.close()
        return "\n\n".join(chunks), None
    except ImportError:
        pass
    except Exception as e:
        return "", f"pymupdf error: {e}"
    # pdfplumber
    try:
        import pdfplumber
        text_chunks = []
        with pdfplumber.open(str(path)) as pdf:
            for p in pdf.pages:
                t = p.extract_text() or ""
                text_chunks.append(t)
        return "\n\n".join(text_chunks), None
    except ImportError:
        pass
    except Exception as e:
        return "", f"pdfplumber error: {e}"
    # PyPDF2
    try:
        import PyPDF2
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            chunks = []
            for p in reader.pages:
                chunks.append(p.extract_text() or "")
        return "\n\n".join(chunks), None
    except ImportError:
        return "", "No PDF library available. Install pymupdf, pdfplumber, or PyPDF2."
    except Exception as e:
        return "", f"PyPDF2 error: {e}"


def _extract_docx(path: Path) -> tuple[str, Optional[str]]:
    try:
        from docx import Document  # python-docx
        doc = Document(str(path))
        paras = [p.text for p in doc.paragraphs if p.text]
        # Tables
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                paras.append(" | ".join(cells))
        return "\n".join(paras), None
    except ImportError:
        return "", "python-docx not installed"
    except Exception as e:
        return "", f"docx error: {e}"


def _extract_xlsx(path: Path) -> tuple[str, Optional[str]]:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        chunks = []
        for sname in wb.sheetnames:
            ws = wb[sname]
            chunks.append(f"## Sheet: {sname}")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= 1000:  # cap rows per sheet
                    chunks.append(f"... ({ws.max_row} rows total, truncated)")
                    break
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    chunks.append(" | ".join(cells))
            chunks.append("")
        return "\n".join(chunks), None
    except ImportError:
        return "", "openpyxl not installed"
    except Exception as e:
        return "", f"xlsx error: {e}"


def _extract_text(path: Path, max_size: int = MAX_TEXT_FILE_SIZE) -> tuple[str, Optional[str]]:
    try:
        size = path.stat().st_size
        if size > max_size:
            return "", f"file too large ({size}B > {max_size}B)"
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read(), None
    except Exception as e:
        return "", f"read error: {e}"


# ============================================================
# Main entry
# ============================================================

def save_upload(webapp_dir: Path, conv_id: str, filename: str, data: bytes,
                mime: Optional[str] = None) -> dict:
    """Save uploaded file + classify + extract content. Ritorna dict descrittore."""
    if not filename:
        return {"error": "filename required"}
    if len(data) > MAX_FILE_SIZE:
        return {"error": f"file too large (max {MAX_FILE_SIZE // 1024 // 1024}MB)"}
    if not mime:
        mime = mimetypes.guess_type(filename)[0] or "application/octet-stream"

    category = classify(filename, mime)
    if category == "image" and len(data) > MAX_IMAGE_SIZE:
        return {"error": f"image too large (max {MAX_IMAGE_SIZE // 1024 // 1024}MB)"}

    # Safe filename
    safe_fname = re.sub(r"[^a-zA-Z0-9._-]+", "_", Path(filename).name)[:200] or "file"
    file_id = secrets.token_hex(6)
    save_dir = conv_uploads_dir(webapp_dir, conv_id)
    full_path = save_dir / f"{file_id}-{safe_fname}"

    try:
        with open(full_path, "wb") as f:
            f.write(data)
    except Exception as e:
        return {"error": f"save failed: {e}"}

    descriptor = {
        "file_id": file_id,
        "filename": filename,
        "saved_filename": full_path.name,
        "category": category,
        "mime": mime,
        "size_bytes": len(data),
        "path": str(full_path),
        "created_at": datetime.utcnow().isoformat() + "Z",
    }

    # Extract content based on category
    extract_err = None
    extracted_text = ""
    image_b64 = None

    if category == "pdf":
        extracted_text, extract_err = _extract_pdf(full_path)
    elif category == "docx":
        extracted_text, extract_err = _extract_docx(full_path)
    elif category == "xlsx":
        extracted_text, extract_err = _extract_xlsx(full_path)
    elif category == "text":
        extracted_text, extract_err = _extract_text(full_path)
    elif category == "image":
        # Encode base64 per vision providers (preview limit 500x500 thumb)
        image_b64 = base64.b64encode(data).decode("ascii")
        descriptor["image_b64_len"] = len(image_b64)
    elif category == "audio":
        # Defer STT to pipeline integration (riusa Whisper Telegram)
        descriptor["needs_stt"] = True
    elif category == "binary":
        extracted_text = f"(binary file, {len(data)} bytes, no extraction available)"

    if extract_err:
        descriptor["extract_error"] = extract_err
    if extracted_text:
        # Cap chars
        if len(extracted_text) > MAX_EXTRACTED_CHARS:
            descriptor["truncated"] = True
            extracted_text = extracted_text[:MAX_EXTRACTED_CHARS] + "\n\n[... truncated ...]"
        descriptor["extracted_text"] = extracted_text
        descriptor["extracted_chars"] = len(extracted_text)
        # Preview = first ~200 chars
        descriptor["preview"] = extracted_text[:200].strip() + ("…" if len(extracted_text) > 200 else "")
    elif image_b64:
        descriptor["image_b64"] = image_b64
        descriptor["preview"] = f"📷 image {Path(filename).suffix.upper()} {len(data)//1024}KB"
    else:
        descriptor["preview"] = f"📎 {Path(filename).suffix.upper()} {len(data)//1024}KB"

    return descriptor


def list_uploads(webapp_dir: Path, conv_id: str) -> list:
    d = conv_uploads_dir(webapp_dir, conv_id)
    items = []
    if not d.is_dir():
        return items
    for f in sorted(d.iterdir(), reverse=True):
        if not f.is_file():
            continue
        items.append({
            "saved_filename": f.name,
            "size_bytes": f.stat().st_size,
            "modified": datetime.fromtimestamp(f.stat().st_mtime).isoformat(),
        })
    return items


def cleanup_old_uploads(webapp_dir: Path, max_age_days: int = 30, keep_conv_ids: Optional[set] = None) -> dict:
    """F24.b — Rimuove file upload più vecchi di max_age_days.

    keep_conv_ids: set di conv_id da preservare anche se vecchi (es. chat ancora aperte).
    Ritorna {removed_files, removed_dirs, freed_bytes}.
    """
    root = uploads_root(webapp_dir)
    if not root.is_dir():
        return {"removed_files": 0, "removed_dirs": 0, "freed_bytes": 0}
    cutoff = datetime.utcnow().timestamp() - (max_age_days * 86400)
    keep = keep_conv_ids or set()
    n_files = 0
    n_dirs = 0
    n_bytes = 0
    for conv_dir in root.iterdir():
        if not conv_dir.is_dir():
            continue
        if conv_dir.name in keep:
            continue
        # Rimuovi file vecchi nella dir
        dir_empty = True
        for f in conv_dir.iterdir():
            if not f.is_file():
                dir_empty = False
                continue
            try:
                if f.stat().st_mtime < cutoff:
                    n_bytes += f.stat().st_size
                    f.unlink()
                    n_files += 1
                else:
                    dir_empty = False
            except Exception:
                dir_empty = False
        # Cleanup dir vuota
        if dir_empty:
            try:
                conv_dir.rmdir()
                n_dirs += 1
            except Exception:
                pass
    return {"removed_files": n_files, "removed_dirs": n_dirs, "freed_bytes": n_bytes}


def delete_upload(webapp_dir: Path, conv_id: str, saved_filename: str) -> bool:
    d = conv_uploads_dir(webapp_dir, conv_id)
    # Safety: enforce within conv dir
    if "/" in saved_filename or ".." in saved_filename:
        return False
    target = d / saved_filename
    if not target.is_file():
        return False
    try:
        target.unlink()
        return True
    except Exception:
        return False


# ============================================================
# Prompt injection: convert descriptors → markdown blocks for user_prompt
# ============================================================

def attachments_to_prompt_block(descriptors: list[dict]) -> str:
    """Per text-based files: produci blocco markdown da concatenare al user prompt.

    Image: NON gestita qui (servono content blocks vision-aware nel provider stream).
    Audio: NON gestita qui (richiede STT pipeline).
    """
    if not descriptors:
        return ""
    blocks = ["\n\n---", "## 📎 Allegati"]
    for d in descriptors:
        cat = d.get("category", "binary")
        fname = d.get("filename", "?")
        if cat == "image":
            blocks.append(f"\n*[Image: `{fname}` allegata — vedi content block vision]*")
            continue
        if cat == "audio":
            transcript = d.get("extracted_text") or ""
            if transcript:
                model = d.get("stt_model", "whisper")
                blocks.append(f"\n### 🎤 {fname} — transcript ({model})")
                blocks.append(f"```\n{transcript}\n```")
            else:
                err = d.get("extract_error") or "no transcript"
                blocks.append(f"\n*[Audio: `{fname}` — {err}]*")
            continue
        text = d.get("extracted_text") or d.get("preview") or ""
        if not text:
            blocks.append(f"\n*[{fname}: contenuto non estratto]*")
            continue
        # Language hint per syntax highlight
        ext = Path(fname).suffix.lower().lstrip(".")
        lang_map = {"py": "python", "js": "javascript", "ts": "typescript", "json": "json",
                    "yaml": "yaml", "yml": "yaml", "md": "markdown", "html": "html", "csv": "csv",
                    "sh": "bash", "sql": "sql"}
        lang = lang_map.get(ext, "")
        blocks.append(f"\n### {fname} ({cat}, {d.get('size_bytes', 0)} B)")
        blocks.append(f"```{lang}")
        blocks.append(text)
        blocks.append("```")
        if d.get("truncated"):
            blocks.append("*(content truncated)*")
    return "\n".join(blocks)


def get_image_attachments(descriptors: list[dict]) -> list[dict]:
    """Filtra solo image descriptors con base64 utile per vision content blocks."""
    return [d for d in (descriptors or []) if d.get("category") == "image" and d.get("image_b64")]
